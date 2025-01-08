"""网约车平台 PDF 格式电子行程单解析器

滴滴出行，高德地图等当代打车软件，都会提供电子行程单导出功能，主要用于财务报销时候的凭据。
但是 PDF 格式不方便系统进行数据的分析和读取，本工具提供了一个简单的命令，用于将 PDF 格式
行程单导出成 CSV 格式或者 Excel 格式。
"""
import os
import sys
import argparse
import re
import logging
import pandas
import tabula

from codecs import BOM_UTF8
from argparse import ArgumentParser
from argparse import ArgumentDefaultsHelpFormatter
from pdfminer.high_level import extract_text
from pdfminer.pdfparser import PDFSyntaxError
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def _parse_didi(file_path, line_count=0):
    """解析滴滴出行的行程单，表头行需要特别清洗"""
    didi_area = [266, 42.765625, 785.028125, 564.134375]
    if line_count != 0:
        didi_area[2] = 3120.003125 + 31.2375 * line_count
    dfs = tabula.read_pdf(file_path, pages="all", area=didi_area)
    if 0 == len(dfs):
        logging.error("no table found")
    elif 1 < len(dfs):
        logging.warning("more than 1 table recognized")
    df = dfs[0]
    # 滴滴的表头行有一些多余的换行符，导致导出的 CSV 破损
    df.columns = [name.replace('\r', ' ') for name in df.columns]
    return df


def _parse_huaxiaozhu(file_path, line_count=0):
    """解析花小猪打车行程单"""
    huaxiaozhu_area = [222, 42, 780, 564]
    if line_count != 0:
        huaxiaozhu_area[2] = 262 + 31 * line_count
    dfs = tabula.read_pdf(file_path, pages="all", area=huaxiaozhu_area)
    if 0 == len(dfs):
        logging.error("no table found")
    elif 1 < len(dfs):
        logging.warning("more than 1 table recognized")
    df = dfs[0]
    # 滴滴的表头行有一些多余的换行符，导致导出的 CSV 破损
    df.columns = [name.replace('\r', ' ') for name in df.columns]
    return df


def _parse_gaode(file_path, line_count=0):
    """解析高德地图的行程单"""
    gaode_area = [173, 37.5767, 791.3437, 559.1864]
    if line_count != 0:
        gaode_area[2] = 216.9033 + line_count * 30
    dfs = tabula.read_pdf(file_path, pages="all", area=gaode_area, stream=True)
    if 0 == len(dfs):
        logging.error("no table found")
    elif 1 < len(dfs):
        logging.warning("more than 1 table recognized")
    return dfs[0]


def _parse_shouqi(file_path, line_count=0):
    """解析首汽约车的行程单，是最难处理的一种类型"""
    shouqi_area = [153.584375, 29.378125, 817.753125, 566.365625]
    if line_count != 0:
        shouqi_area[2] = 176.64062 + 15.95379 * line_count
    dfs = tabula.read_pdf(file_path, pages="all",
                          area=shouqi_area, stream=True)
    if 0 == len(dfs):
        logging.error("no table found")
    elif 1 < len(dfs):
        logging.warning("more than 1 table recognized")
    df = dfs[0]

    # 对识别结果进行处理
    # 表头处理
    rows = df.iloc[0].values
    df.columns = [str(x).strip() + ('' if str(y).strip() == 'nan' else str(y).strip()) for x, y in
                  zip(df.columns, rows)]

    # 数据处理
    data = df.values
    row_index = range(len(data))
    new_data = []
    for x, y in zip(row_index[1::2], row_index[2::2]):
        new_row = [str(a).strip() + ('' if str(b).strip() == 'nan' else str(b).strip()) for a, b in
                   zip(data[x], data[y])]
        new_data.append(new_row)

    new_df = pandas.DataFrame(new_data, columns=df.columns)
    return new_df


def _parse_meituan(file_path, line_count=0):
    """解析美团打车的行程单，也是比较难处理的一种类型"""
    meituan_area = [285.7275, 41.6925, 314.7975, 571.0725]
    if line_count:
        meituan_area[2] = 314.7975 + 28.305 * line_count
    dfs = tabula.read_pdf(file_path, pages='1', area=meituan_area, stream=True)
    if 0 == len(dfs):
        logging.error("no table found")
    elif 1 < len(dfs):
        logging.warning("more than 1 table recognized")

    df = dfs[0]

    data = df.values
    row_index = range(len(data))
    new_data = []
    for x, y in zip(row_index[::2], row_index[1::2]):
        new_row = [('' if str(x).strip() == 'nan' else (str(x).strip() + ' ')) + str(y).strip() for x, y in
                   zip(data[x], data[y])]
        new_data.append(new_row)
    new_df = pandas.DataFrame(new_data, columns=df.columns)

    return new_df


def _parse_unknown(file_path):
    dfs = tabula.read_pdf(file_path, pages="all", stream=True)
    if 0 == len(dfs):
        logging.error("no table found")
    elif 1 < len(dfs):
        logging.warning("more than 1 table recognized")
    return dfs[0]


def _output_csv(df, output_path):
    """利用 DataFrame 自身的 API，导出到 CSV 格式"""
    # 增加 BOM 头，否则不能双击Excel 直接打开CSV
    with open(output_path, mode='wb') as output:
        output.write(BOM_UTF8)

    with open(output_path, mode='a', newline='') as output:
        df.to_csv(output, index=False)


def _output_excel(df, output_path):
    """利用 DataFrame 自身的 API，导出到 Excel 格式"""
    df.to_excel(output_path, index=False, sheet_name='Sheet1')


def _output(df, file_type):
    extension = {'csv': 'csv', 'excel': 'xlsx'}
    if file_type in ['csv', 'excel']:
        exporter = getattr(sys.modules[__name__], '_output_' + file_type)
        exporter(df, 'output.' + extension[file_type])
    else:
        logging.error('不支持的导出文件类型，目前仅支持 CSV 和 Excel')


platform_pattern = {
    'didi':    {'title_like': '滴滴出行', 'line_count_like': r'共(\d+)笔行程', 'parser': _parse_didi},
    'gaode':   {'title_like': '高德地图', 'line_count_like': r'共计(\d+)单行程', 'parser': _parse_gaode},
    'shouqi':  {'title_like': '首汽约车电子行程单', 'line_count_like': r'共(\d+)个行程', 'parser': _parse_shouqi},
    'meituan': {'title_like': '美团打车', 'line_count_like': r'(\d+)笔行程', 'parser': _parse_meituan},
    'huaxiaozhu': {'title_like': '花小猪打车', 'line_count_like': r'(\d+)笔行程', 'parser': _parse_huaxiaozhu}
}


def _extract_text(file_path):
    try:
        pdf_to_text = extract_text(file_path)
        return ''.join([x for x in filter(lambda x: x.strip() != '', "".join(pdf_to_text).splitlines())])
    except PDFSyntaxError as pse:
        logging.error('文件解析错误，不是一个正确的 PDF 文件')
        raise Exception('无法解析 PDF') from pse
    return ''


def _read_meta(file_path):
    """读取行程单的信息，识别平台、行数、页数等"""
    file_content = _extract_text(file_path)
    line_count = 0
    for p, pattern in platform_pattern.items():
        if re.search(pattern['title_like'], file_content):
            match = re.search(pattern['line_count_like'], file_content)
            if match:
                line_count = int(match.group(1))
            return p, line_count, pattern['parser']
    return 'unknown', 0, _parse_unknown


def main(args=None):
    arg_parser = ArgumentParser(description=__doc__, add_help=True,
                                formatter_class=ArgumentDefaultsHelpFormatter)
    arg_parser.add_argument('file_path', metavar='<FILE>', type=str,
                            help='需要处理的行程单文件')
    arg_parser.add_argument('--output_type', '-t', metavar='<TYPE>', type=str, default='csv',
                            help='输出文件类型，默认是csv，也可以是excel')
    args = arg_parser.parse_args(args)

    platform, line_count, parser = _read_meta(args.file_path)
    print("识别出来的平台是：%s，行程行数是：%d" % (platform, line_count))

    df = parser(args.file_path, line_count)
    _output(df, args.output_type)

    print(df)




# 创建报告的函数，接受填表人、填表日期和金额合计作为参数
def generate_report(fill_person, fill_date, data, total_amount):
    # 创建 Word 文档
    doc = Document()

    # 添加标题
    title = doc.add_heading(level=1)  # 添加标题
    title_run = title.add_run("市内交通费报销明细表")  # 设置标题内容

    # 设置字体为微软雅黑
    title_run.font.name = '微软雅黑'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 设置字体加粗
    title_run.font.bold = True

    # 设置字体大小为20
    title_run.font.size = Pt(20)

    # 设置字体颜色为黑色（RGB(0, 0, 0)）
    title_run.font.color.rgb = RGBColor(0, 0, 0)

    # 设置标题居中
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 设置标题的行距为单倍行距
    title.paragraph_format.line_spacing = 1.0  # 单倍行距

    # 添加填表人信息
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(f"填表人：{fill_person}\t\t\t\t\t\t填表日期：{fill_date}")

    # 设置字体为宋体
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 设置字体大小为小四（即 12pt）
    run.font.size = Pt(12)

    # 设置单倍行距
    paragraph.paragraph_format.line_spacing = 1.0  # 单倍行距

    # 添加表格
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'

    # 表头
    headers = ["序号", "乘车日期", "乘车人", "出发地", "目的地", "外出事由", "金额"]
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.name = '宋体'
        hdr_cells[i].paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(12)

    # 添加表格内容（所有行都设置为宋体、小四、单倍行距）
    for row_data in data:
        row_data = [str(item) for item in row_data]
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            cell = row_cells[i]
            cell.text = cell_text
            cell.paragraphs[0].runs[0].font.name = '宋体'
            cell.paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            cell.paragraphs[0].runs[0].font.size = Pt(12)
        # 设置行的单倍行距
        row_cells[0].paragraphs[0].paragraph_format.line_spacing = 1.0

    # 添加金额合计（占四列，金额占三列），并将金额作为参数传递
    total_row = table.add_row()
    total_row.cells[0].merge(total_row.cells[3])  # 合并前四列，显示“金额合计”
    total_row.cells[0].text = "金额合计："
    total_row.cells[0].paragraphs[0].runs[0].font.name = '宋体'
    total_row.cells[0].paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    total_row.cells[0].paragraphs[0].runs[0].font.size = Pt(12)
    total_row.cells[0].paragraphs[0].runs[0].font.bold = True  # 设置加粗

    total_row.cells[4].merge(total_row.cells[6])  # 合并后面三列，显示金额数字
    total_row.cells[4].text = str(total_amount)  # 使用参数传递金额
    total_row.cells[4].paragraphs[0].runs[0].font.name = '宋体'
    total_row.cells[4].paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    total_row.cells[4].paragraphs[0].runs[0].font.size = Pt(12)
    total_row.cells[4].paragraphs[0].runs[0].font.bold = True  # 设置加粗

    # 添加备注（特殊事项说明）
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("特殊事项说明：")
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(11)
    run.font.bold = True  # 设置加粗
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 左对齐
    paragraph.paragraph_format.line_spacing = 1.0  # 单倍行距

    # 保存文档
    doc.save("市内交通费报销明细表.docx")
if __name__ == '__main__':
    # sys.exit(main())
    # 手动设置参数
    file_path = r'滴滴出行行程报销单.pdf'  # 设置需要处理的PDF文件路径
    passenger = "passenger"  # 乘车人
    year = "2024"  # 乘车年份
    date = "2025.01.08"  #  制表日期
    # 配置常用外出事由
    configs = {
        ('出发地',  '目的地', "参加XXX会议"),
        ('YYY', 'XXX', '参加XXX会议-返回'),
    }
    platform, line_count, parser = _read_meta(file_path)

    df = parser(file_path, line_count)
    # 转换为列表
    data_list = df.values.tolist()
    # 提取特定元素
    extracted_elements = [[row[0], year+"-"+str(row[2]).split(" ")[0], passenger, str(row[4]).split("|")[-1].replace("-", ""), str(row[5]).split("|")[-1].replace("-", ""), '', row[7]] for row in data_list]

    total_amount = sum([row[7] for row in data_list])


    data = []
    for element in extracted_elements:
        from_, to_ = element[3], element[4]
        for cfg in configs:
            if from_.find(cfg[0])!=-1 and to_.find(cfg[1])!=-1:
                element[5] = cfg[2]
                data.append(element)
                break



    generate_report(passenger, date, data, total_amount)